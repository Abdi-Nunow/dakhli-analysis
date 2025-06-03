import streamlit as st
import pandas as pd
import sqlite3
from datetime import datetime
import io
from docx import Document
import plotly.express as px

st.set_page_config(page_title="Dakhli Analysis App", layout="wide")

# Background color
page_bg_color = """
<style>
    .stApp {
        background-color: #f0f0f5;
    }
</style>
"""
st.markdown(page_bg_color, unsafe_allow_html=True)

st.title("Falanqaynta Dakhliga Maalinlaha ah ee Canshuuraha")

zones = {
    "Afder": ["Bare", "Elekere", "GodGod", "Hargelle", "Mirab Imi", "Ilig Dheere", "Raaso", "Qooxle", "Doollo bay", "Baarey", "Washaaqo", "Ciid Laami", "Xagar Moqor"],
    "Dhawa": ["Hudet", "Lahey", "Mubaarak", "Qadhaadhumo", "Malka Mari", "Ceel Goof", "Ceel Orba", "Dheer Dheertu", "Ceel Dheer"],
    "Dollo": ["Boh", "Danot", "Daratole", "Geladin", "Gal-Hamur", "Lehel-Yucub", "Warder", "Yamarugley", "Urmadag"],
    "Erer": ["Fiq", "Lagahida", "Mayaa-muluqo", "Qubi", "Salahad", "Waangaay", "Xamaro", "Yaxoob"],
    "Fafan": ["Awbare", "Babille", "Goljano", "Gursum", "Harawo", "Haroorays", "Harshin", "Jijiga", "Kebri Beyah", "Qooraan", "Shabeeley", "Wajale", "Tuli Guled"],
    "Jarar": ["Araarso", "Awaare", "Bilcil Buur", "Birqod", "Daroor", "Degehabur", "Dhagaxmadow", "Dig", "Gunagado", "Misraq Gashamo", "Yoocaale"],
    "Korahe": ["Boodaley", "Ceel-Ogadeen", "Dobawein", "Higloley", "Kebri Dahar", "Kudunbuur", "Laas-dhankayre", "Marsin", "Shekosh", "Shilavo"],
    "Liben": ["Bokolmayo", "Deka Softi", "Dollo Ado", "Filtu", "Kersa Dula", "Gooro Bakaksa", "Gurra Damole"],
    "Nogob": ["Ayun", "Duhun", "Elweyne", "Gerbo", "Hararey", "Hora-shagax", "Segeg"],
    "Shabelle": ["Abaaqoorow", "Adadle", "Beercaano", "Danan", "Elele", "Ferfer", "Gode", "Imiberi", "Kelafo", "Mustahil"],
    "Sitti": ["Adigala", "Afdem", "Ayesha", "Bike", "Dambal", "Erer", "Gablalu", "Mieso", "Shinile", "Dhunyar", "Daymeed"]
}

kable_list = [f"Kable {i:02}" for i in range(1, 41)]
tax_types = ["VAT", "TOT", "INCOME TAX", "PROFIT TAX", "LAND TAX", "PROPERTY TAX", "EXERCISE TAX", "OTHER TAX"]
years = [str(year) for year in range(1990, datetime.now().year + 10)]

# Connect to SQLite DB (file created if not exist)
conn = sqlite3.connect('dakhli_data.db', check_same_thread=False)
c = conn.cursor()

# Create table if not exists
c.execute('''
CREATE TABLE IF NOT EXISTS dakhli (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    Magaca TEXT,
    TIN TEXT,
    Mobile TEXT,
    Zone TEXT,
    District TEXT,
    Kable TEXT,
    Tax_Type TEXT,
    Tax_Year TEXT,
    Date TEXT,
    Income REAL,
    Payment REAL,
    Outstanding REAL
)
''')
conn.commit()

st.header("Geli Xogta Canshuur Bixiyaha")

col1, col2, col3 = st.columns(3)
with col1:
    name = st.text_input("Magaca Canshuur Bixiyaha")
    mobile = st.text_input("Mobile")
    zone = st.selectbox("Gobolka (Zone)", list(zones.keys()))
with col2:
    tin = st.text_input("TIN No (10)", max_chars=10)
    if tin and (not tin.isdigit() or len(tin) != 10):
        st.error("TIN waa inuu noqdaa 10 lambar")
    district = st.selectbox("Degmada (District)", zones[zone])
    kable = st.selectbox("Kable", kable_list)
with col3:
    tax_type = st.selectbox("Nooca Canshuurta", tax_types)
    tax_year = st.selectbox("Sanadka Canshuurta", years)
    date = st.date_input("Taariikhda", datetime.today())

income = st.number_input("Dakhliga Lagu Leeyahay", min_value=0.0, step=0.1)
payment = st.number_input("Lacagta la Bixiyay", min_value=0.0, step=0.1)

if st.button("Kaydi Xogta"):
    if not tin.isdigit() or len(tin) != 10:
        st.error("Fadlan geli TIN sax ah (10 lambar)")
    elif not name.strip():
        st.error("Fadlan geli magaca canshuur bixiyaha")
    else:
        outstanding = income - payment
        c.execute('''
            INSERT INTO dakhli (Magaca, TIN, Mobile, Zone, District, Kable, Tax_Type, Tax_Year, Date, Income, Payment, Outstanding)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (name, tin, mobile, zone, district, kable, tax_type, tax_year, date.strftime('%Y-%m-%d'), income, payment, outstanding))
        conn.commit()
        st.success("Xogta waa la keydiyay ✅")

st.header("Falanqaynta Dakhliga")

df = pd.read_sql_query("SELECT * FROM dakhli", conn)

if df.empty:
    st.info("Xog lama helin. Fadlan geli xog si aad u aragto falanqayn.")
else:
    # Display table
    st.dataframe(df)

    today = datetime.today().strftime('%Y-%m-%d')
    total_clients = df[df['Date'] == today].shape[0]
    total_income = df['Income'].sum()
    total_outstanding = df['Outstanding'].sum()

    st.metric("Tirada Canshuur Bixiyayaasha Maanta", total_clients)
    st.metric("Dakhliga Guud", f"{total_income:,.2f}")
    st.metric("Lacagta aan wali la Bixin", f"{total_outstanding:,.2f}")

    fig = px.bar(df, x='Date', y='Income', color='Tax_Type', title='Dakhliga Maalinlaha ah')
    st.plotly_chart(fig, use_container_width=True)

    # Excel download
    excel_buffer = io.BytesIO()
    df.to_excel(excel_buffer, index=False, engine='xlsxwriter')
    st.download_button(
        label="📥 Soo Degso Excel",
        data=excel_buffer.getvalue(),
        file_name="dakhli_data.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    # Word download
    doc = Document()
    doc.add_heading("Xogta Dakhliga Canshuur Bixiyayaasha", 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    word_buffer = io.BytesIO()
    doc.save(word_buffer)

    st.download_button(
        label="📥 Soo Degso Word",
        data=word_buffer.getvalue(),
        file_name="dakhli_data.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
